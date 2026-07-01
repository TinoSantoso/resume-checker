"""Generate DOCX fixtures for tests.

Run from project root:

    .venv/bin/python scripts/gen_docx_fixtures.py

Produces a `data/tino_shaped.docx` that mirrors the content of
`tino_actual.pdf` so we have a real, hand-built DOCX to assert against.

Also produces a `data/strong_candidate.docx` (strong CV in DOCX form)
so the existing TestScoreCv::test_strong_sample_scores_high coverage
extends to DOCX.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def build_tino_shaped(out: Path) -> None:
    """Mirror data/tino_actual.pdf content as DOCX."""
    d = Document()
    d.add_heading("TINO APRIKA SANTOSO", 0)
    d.add_paragraph("Full Stack Developer")
    d.add_paragraph(
        "tino.santoso92@gmail.com | 6285186682433 | South Tangerang"
    )
    d.add_paragraph(
        "Passionate Full Stack developer in building comprehensive web applications "
        "using Laravel and CodeIgniter. Expertise in integrating APIs, particularly "
        "payment gateways, to enhance functionality and user experience. I create "
        "dynamic and responsive front-ends with React and develop practical desktop "
        "applications with Python and TkInter, exemplified by work on relay turnstile systems."
    )

    d.add_heading("Experience", level=1)
    d.add_paragraph(
        "Fullstack Developer  Essity - GOP 6, Serpong Indonesia  June 2025 - August 2025"
    )
    d.add_paragraph(
        "As a Fullstack Developer, developed and maintained comprehensive CRM "
        "solutions to optimize sales operations and client management. Built systems "
        "for tracking sales team visits and integrated payment gateway APIs."
    )
    d.add_paragraph(
        "Fullstack Developer  PT. Yapindo  June 2020 - June 2025"
    )
    d.add_paragraph(
        "Developed Laravel-based web applications, optimized MySQL queries "
        "(40% faster), and mentored 2 junior developers."
    )

    d.add_heading("Skills", level=1)
    d.add_paragraph(
        "PHP - Advanced, Javascript - Advanced, Python - Intermediate, "
        "Laravel - Advanced, CodeIgniter - Advanced, React - Intermediate, "
        "MySQL - Advanced, PostgreSQL - Advanced, GIT - Advanced"
    )
    d.add_paragraph(
        "Languages: English - Intermediate, Indonesia - Native"
    )

    d.add_heading("Education", level=1)
    d.add_paragraph(
        "Bachelor's Degree in Informatics  Universitas Islam Indonesia - Yogyakarta"
    )
    d.add_paragraph("GPA: 2.86")

    d.add_heading("Links", level=1)
    d.add_paragraph(
        "Linkedin: linkedin.com/in/tino-santoso/  Github: github.com/tinosantoso"
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)


def build_strong_candidate(out: Path) -> None:
    """A strong, well-structured CV as DOCX — used to verify that the
    DOCX path produces high scores just like the PDF strong sample."""
    d = Document()
    d.add_heading("JANE STRONG", 0)
    d.add_paragraph("jane.strong@example.com | +1 555 010 2024 | San Francisco, CA")
    d.add_paragraph(
        "linkedin.com/in/janestrong | github.com/janestrong"
    )

    d.add_heading("Professional Summary", level=1)
    d.add_paragraph(
        "Senior Software Engineer with 8 years of experience designing and "
        "shipping distributed systems at scale. Led teams of up to 5 engineers, "
        "owned service reliability for revenue-critical paths (99.99% SLO), and "
        "drove a 35% latency reduction via async processing. Passionate about "
        "clean architecture, observability, and mentoring."
    )

    d.add_heading("Experience", level=1)
    d.add_paragraph("Senior Software Engineer  BigTech Co  Jan 2021 - Present")
    d.add_paragraph(
        "Architected event-driven pipeline processing 50K events/second.",
        style="List Bullet",
    )
    d.add_paragraph(
        "Led migration to Kubernetes reducing infrastructure cost by 30%.",
        style="List Bullet",
    )
    d.add_paragraph(
        "Mentored 4 junior engineers; 2 promoted to Senior within a year.",
        style="List Bullet",
    )
    d.add_paragraph("Software Engineer  StartupX  Jun 2017 - Dec 2020")
    d.add_paragraph(
        "Built payment integration handling $120M/year.",
        style="List Bullet",
    )
    d.add_paragraph(
        "Reduced API p99 latency from 800ms to 220ms via caching + query optimization.",
        style="List Bullet",
    )
    d.add_paragraph(
        "Implemented automated deployment pipeline cutting release time from 4 hours to 12 minutes.",
        style="List Bullet",
    )

    d.add_heading("Skills", level=1)
    d.add_paragraph(
        "Python, Go, TypeScript, React, PostgreSQL, Redis, Kafka, "
        "AWS, Kubernetes, Docker, Terraform, gRPC, GraphQL, Prometheus"
    )

    d.add_heading("Education", level=1)
    d.add_paragraph(
        "M.S. Computer Science  Stanford University  2017"
    )
    d.add_paragraph(
        "B.S. Computer Science  UC Berkeley  2015  GPA: 3.9"
    )

    d.add_heading("Certifications", level=1)
    d.add_paragraph("AWS Solutions Architect Professional  2023")
    d.add_paragraph("Certified Kubernetes Administrator  2022")

    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)


def main() -> int:
    data_dir = Path(__file__).resolve().parent.parent / "data"
    fixtures = [
        (data_dir / "tino_shaped.docx", build_tino_shaped),
        (data_dir / "strong_candidate.docx", build_strong_candidate),
    ]
    for out, builder in fixtures:
        builder(out)
        size = out.stat().st_size
        print(f"wrote {out} ({size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
