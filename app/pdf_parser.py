"""CV parser — extract raw text + naive section segmentation.

Supports PDF (via PyMuPDF, with layout-aware column handling) and DOCX
(via python-docx). Both formats are normalized to the same
``"\\n".join(line)`` shape so downstream section segmentation works
identically.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Union

import pymupdf  # PyMuPDF

from .redactor import redact, unmask


# --------------------------------------------------------------------------- #
# Parsed result container
# --------------------------------------------------------------------------- #

@dataclass
class ParsedCV:
    """Result of parsing a CV file.

    Attributes:
        sections: dict of section_name -> section_text. PII has been
            redacted (replaced with placeholders like ``[EMAIL_1]``).
        pii_map: mapping from placeholder back to original value, so the
            recruiter-facing UI can ``unmask()`` before display.

    Iterating over a ``ParsedCV`` yields ``(section_name, text)`` tuples
    for backward compatibility with code that treated the parse result
    as a dict.
    """

    sections: Dict[str, str] = field(default_factory=dict)
    pii_map: Dict[str, str] = field(default_factory=dict)
    # Raw (unredacted) Header + Contact + Links text. Used by Contact
    # scoring only — see ``app.scorer._score_contact``. Empty when
    # ``redact_pii=False`` (nothing to unmask) or when those sections
    # are absent from the parsed doc.
    raw_contact: str = ""

    def __iter__(self):
        return iter(self.sections.items())

    def __getitem__(self, key: str) -> str:
        return self.sections[key]

    def __setitem__(self, key: str, value: str) -> None:
        """Allow score-time mutation (e.g. carving Summary out of Header)."""
        self.sections[key] = value

    def __contains__(self, key: str) -> bool:
        return key in self.sections

    def get(self, key: str, default: str = "") -> str:
        """dict.get() compatibility — used by scorer.py."""
        return self.sections.get(key, default)

    def __len__(self) -> int:
        return len(self.sections)

    def keys(self):
        return self.sections.keys()

    def values(self):
        return self.sections.values()

    def items(self):
        return self.sections.items()

    def render_unmasked(self) -> Dict[str, str]:
        """Return sections with original PII restored (for recruiter UI)."""
        return {k: unmask(v, self.pii_map) for k, v in self.sections.items()}

# Heuristic section headers (case-insensitive). Order matters for priority.
SECTION_PATTERNS = {
    "Contact": [
        r"^\s*contact(\s+info(rmation)?)?\s*$",
        r"^\s*personal\s+(info(rmation)?|details)\s*$",
    ],
    "Summary": [
        r"^\s*(professional\s+)?summary\s*$",
        r"^\s*profile\s*$",
        r"^\s*objective\s*$",
        r"^\s*about(\s+me)?\s*$",
    ],
    "Experience": [
        r"^\s*(work\s+)?experience\s*$",
        r"^\s*employment(\s+history)?\s*$",
        r"^\s*professional\s+experience\s*$",
        r"^\s*career\s+history\s*$",
    ],
    "Skills": [
        r"^\s*(technical\s+)?skills\s*$",
        r"^\s*technologies\s*$",
        r"^\s*tech(nical)?\s+stack\s*$",
    ],
    "Education": [
        r"^\s*education(al)?(\s+background)?\s*$",
        r"^\s*academic(\s+background)?\s*$",
    ],
    "Projects": [
        r"^\s*(side\s+)?projects\s*$",
        r"^\s*key\s+projects\s*$",
    ],
    "Certifications": [
        r"^\s*certifications?\s*$",                                # certification / certifications
        r"^\s*licenses?(\s+(&|and)\s+certifications?)?\s*$",       # license / licenses, optionally + certifications
        r"^\s*certifications?\s+(&|and)\s+licenses?\s*$",
    ],
    "Links": [
        r"^\s*(social\s+media|links?|profiles?|online\s+presence)\s*$",
    ],
}


# --------------------------------------------------------------------------- #
# Extractors
# --------------------------------------------------------------------------- #

def extract_text(pdf_path: Union[str, Path]) -> str:
    """Pull all text from a PDF, one paragraph per line.

    Uses layout-aware extraction (app.layout) so 2-column PDFs come out in
    proper reading order. Single-column PDFs produce identical output to
    PyMuPDF's native ``get_text("text")`` for the spans we keep, but the
    layout module is the single source of truth for both.
    """
    from .layout import extract_text_with_layout
    return extract_text_with_layout(pdf_path)


def _extract_docx_text(docx_path: Union[str, Path]) -> str:
    """Pull text from a DOCX, one paragraph per line.

    Uses python-docx. Tables are rendered row-by-row (cells joined with
    " | " so a 2-column layout doesn't smear together). Empty paragraphs
    are dropped, matching the PDF extractor's behavior.

    Bullet-style paragraphs (``List Bullet``, ``List Number``, etc.) are
    prefixed with a ``• `` marker so the downstream experience heuristic
    (which looks for ``^\\s*([-*•●▪]|\\d+[.)])``) can detect them. Without
    this, a real Word-style bulleted experience section would be flagged
    as "paragraph format" and penalized.
    """
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - import guard
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        ) from e

    document = docx.Document(str(docx_path))
    lines: list[str] = []

    # python-docx paragraph style names that should be rendered as bullets.
    # Match by case-insensitive substring so locale variants ("Daftar Bullet")
    # are also caught when Word is set to Indonesian.
    _BULLET_STYLE_HINTS = ("list bullet", "list paragraph", "daftar")

    def _is_bullet_style(p) -> bool:
        try:
            name = (p.style.name or "").lower()
        except Exception:
            return False
        return any(hint in name for hint in _BULLET_STYLE_HINTS)

    def _is_numbered_style(p) -> bool:
        try:
            name = (p.style.name or "").lower()
        except Exception:
            return False
        return "list number" in name

    # Walk the document body in order — paragraphs and tables alternate.
    # We use the underlying body element to preserve order.
    from docx.oxml.ns import qn

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            # Paragraph — find the matching docx.Paragraph object.
            for p in document.paragraphs:
                if p._element is child:
                    text = p.text.strip()
                    if not text:
                        break
                    if _is_bullet_style(p):
                        lines.append(f"• {text}")
                    elif _is_numbered_style(p):
                        # Numbered list — leave as-is; the heuristic's
                        # `\\d+[.)]` pattern will catch it downstream.
                        lines.append(text)
                    else:
                        lines.append(text)
                    break
        elif tag == qn("w:tbl"):
            # Table — join cells with " | " per row.
            for table in document.tables:
                if table._element is child:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        # Drop empty cells but keep the joiner.
                        non_empty = [c for c in cells if c]
                        if non_empty:
                            lines.append(" | ".join(non_empty))
                    break

    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Section segmentation (format-agnostic; operates on raw text)
# --------------------------------------------------------------------------- #

def _classify_line(line: str) -> str | None:
    """Return section name if line looks like a section header, else None.

    Supports both English and Indonesian section headers (see
    ``app.i18n.SECTION_PATTERNS_ID``).
    """
    low = line.lower().strip()
    for section, patterns in SECTION_PATTERNS.items():
        for pat in patterns:
            if re.match(pat, low):
                return section
    # Indonesian section headers (P1.4)
    from .i18n import SECTION_PATTERNS_ID
    for section, patterns in SECTION_PATTERNS_ID.items():
        for pat in patterns:
            if re.match(pat, low):
                return section
    return None


def segment_sections(raw_text: str) -> Dict[str, str]:
    """Walk text line-by-line, bucket content into detected sections.

    Anything before the first detected header goes into 'Header' (usually
    contact info / name block).
    """
    sections: Dict[str, list[str]] = {
        "Header": [],
        "Contact": [],
        "Summary": [],
        "Experience": [],
        "Skills": [],
        "Education": [],
        "Projects": [],
        "Certifications": [],
        "Links": [],
        "Other": [],
    }
    current = "Header"

    for line in raw_text.splitlines():
        sec = _classify_line(line)
        if sec:
            current = sec
            continue
        sections[current].append(line)

    # Post-pass (audit fix-now): some PDFs put the Summary paragraph
    # physically AFTER another section header (typically Education), so
    # the line-by-line walk above strands it in that other section and
    # only the trailing wrapped line ends up in Summary. If Summary is
    # sparse and we find a prose block elsewhere that looks like Summary
    # content, move it back to Summary.
    _reclaim_stranded_summary(sections)

    # Post-pass: collapsed section headers (e.g. "Skills\nEducation" with
    # nothing between) leave Skills empty and dump skills content into the
    # next section. If Skills is empty and the *following* section in the
    # walk order holds skills-like content, move it back.
    _reclaim_skills_content(sections)

    # Drop empty sections, return as dict[str, str]
    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


# Patterns that identify a line as clearly NOT Summary prose.
_BULLET_RE = re.compile(r"^\s*[•●▪\-\*]\s")
# Date ranges to exclude: "2020 - 2025", "2020–present", "June 2025 - August 2025",
# "Jan 2020 - Dec 2022". Covers both 4-digit-year-only and Month-YYYY variants.
_MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)(?:[a-z]+)?"
_DATE_RANGE_RE = re.compile(
    rf"(?:(?:\b(19|20)\d{2})\s*[\-–—]\s*(?:\d{{4}}|present|now))"
    rf"|(?:(?:\b{_MONTH}\b\.?\s+(?:19|20)\d{{2}})\s*[\-–—]\s*(?:(?:\b{_MONTH}\b\.?\s+(?:19|20)\d{{2}})|present|now))",
    re.I,
)


def _looks_like_summary_line(line: str) -> bool:
    """True if the line reads like Summary prose — long, no bullet, no
    date range, no skill-list comma pattern. Used by the post-pass in
    ``segment_sections`` to recover a Summary paragraph that was stranded
    in a later section by PDF reading order.
    """
    s = line.strip()
    if len(s) < 30:
        return False
    if _BULLET_RE.match(s):
        return False
    if _DATE_RANGE_RE.search(s):
        return False
    # Skill-list pattern: many short comma-separated tokens (≥4 tokens
    # each <18 chars). Real Summary mentions skills too, so this is a
    # weak signal.
    if len(s.split(",")) >= 4 and all(len(t.strip()) < 18 for t in s.split(",")[1:]):
        return False
    return True


def _reclaim_stranded_summary(sections: Dict[str, list[str]]) -> None:
    """Move a contiguous prose block that looks like Summary from another
    section back into the Summary bucket. Mutates ``sections`` in place.
    No-op if Summary already has substantial content (≥3 non-empty lines).
    """
    summary = sections["Summary"]
    # Only reclaim if Summary is clearly truncated: <100 chars total, or empty.
    # The wrapped-paragraph bug dumps only a single ~30-char trailing line into
    # Summary, so a short Summary IS the signal to try reclaiming from later
    # sections. A long Summary means real content already landed there.
    if sum(len(ln.strip()) for ln in summary) >= 100:
        return  # Summary already has substantial content; leave it alone.

    # Only scan sections that may legitimately contain a stranded Summary
    # paragraph due to PDF reading-order quirks: Experience / Education /
    # Other. Header is intentionally excluded — the carve-out helper handles
    # Summary-in-Header at score_cv time, not here. Scanning Header would
    # steal contact info or name lines on DOCX files where the "summary"
    # prose sits inside the Header bucket by design.
    for sec_name in ("Education", "Experience", "Other"):
        lines = sections[sec_name]
        # Find the longest contiguous run of Summary-like lines.
        best_start, best_len = -1, 0
        i = 0
        while i < len(lines):
            if _looks_like_summary_line(lines[i]):
                j = i
                while j < len(lines) and _looks_like_summary_line(lines[j]):
                    j += 1
                run_len = j - i
                if run_len > best_len:
                    best_start, best_len = i, run_len
                i = j
            else:
                i += 1
        if best_len >= 2:
            reclaimed = lines[best_start:best_start + best_len]
            sections["Summary"].extend(reclaimed)
            del lines[best_start:best_start + best_len]
            return  # one block is enough; don't over-reclaim.


def _looks_like_skills_block(lines: list[str]) -> bool:
    """True if the lines read like a Skills list: many comma-separated
    tokens (≥4 distinct items) OR many canonical-skill dictionary hits.

    Used by ``_reclaim_skills_content`` to identify content that was
    stranded in the wrong section by collapsed section headers.
    """
    text = " ".join(lines)
    # Comma-separated token run (≥5 tokens, ≥4 of them short <20 chars).
    tokens = [t.strip() for t in text.split(",") if t.strip()]
    if len(tokens) >= 5 and sum(1 for t in tokens if len(t) < 20) >= 4:
        return True
    # Canonical-skill dictionary hits (lazy import — skill_dictionary is
    # only needed when this post-pass actually runs, not on every parse).
    try:
        from .skill_dictionary import extract_skills
    except Exception:
        return False
    canonical = extract_skills(text)
    return len(canonical) >= 3


# Walk order matches the canonical section order defined in
# ``segment_sections``; we only reclaim into sections that come BEFORE the
# one holding the candidate content (collapsed header = empty Skills sits
# before non-empty Education in the walk order).
_SKILLS_RECLAIM_ORDER = (
    "Experience", "Education", "Other",
)


def _reclaim_skills_content(sections: Dict[str, list[str]]) -> None:
    """If Skills ended up empty (collapsed header swallowed its content),
    scan later sections for a Skills-shaped block and move it back.

    Mirrors ``_reclaim_stranded_summary`` but for Skills. Only scans
    sections that legitimately follow Skills in the canonical walk order,
    so we don't accidentally steal content from earlier sections like
    Experience that happens to mention tools.
    """
    if sections["Skills"]:
        return  # Skills has content; nothing to reclaim.
    for sec_name in _SKILLS_RECLAIM_ORDER:
        lines = sections.get(sec_name, [])
        if not lines:
            continue
        if _looks_like_skills_block(lines):
            sections["Skills"] = lines
            sections[sec_name] = []
            return  # one block is enough.


# --------------------------------------------------------------------------- #
# Public dispatcher
# --------------------------------------------------------------------------- #

_SUPPORTED_SUFFIXES = {".pdf", ".docx"}


def parse_cv(
    cv_path: Union[str, Path],
    *,
    redact_pii: bool = True,
) -> ParsedCV:
    """One-shot helper: extract + segment. Supports PDF and DOCX.

    Args:
        cv_path: path to ``.pdf`` or ``.docx`` file.
        redact_pii: when True (default), PII (email, phone, LinkedIn,
            street address, DOB) is replaced with stable placeholders
            and a mapping is kept so the recruiter UI can ``unmask()``
            before display. Set False only for debugging — production
            code should leave redaction on.

    Returns:
        ``ParsedCV`` with ``.sections`` (redacted) and ``.pii_map``.
        Iterating yields ``(section_name, text)`` tuples so callers that
        treated the old dict return value as iterable still work.
    """
    path = Path(cv_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text(path)
    elif suffix == ".docx":
        text = _extract_docx_text(path)
    else:
        raise ValueError(
            f"Unsupported CV format '{suffix}'. Supported: {sorted(_SUPPORTED_SUFFIXES)}"
        )

    sections = segment_sections(text)

    if not redact_pii:
        # Return a ParsedCV with empty map (no unmask possible). raw_contact
        # is the same as the unredacted Header+Contact+Links+Summary since
        # nothing was stripped — scorer can use either.
        raw_contact = "\n".join(
            sections.get(name, "") for name in ("Header", "Contact", "Links", "Summary")
        )
        return ParsedCV(sections=sections, pii_map={}, raw_contact=raw_contact)

    # Redact per-section so the mapping accumulates correctly across the
    # whole document (placeholders stay globally unique, not per-section).
    full_redacted = "\n".join(
        f"{name}\n{body}" for name, body in sections.items()
    )
    redacted_text, pii_map = redact(full_redacted)

    # Re-segment after redaction. The section headers themselves don't
    # contain PII so the boundaries are preserved.
    redacted_sections = segment_sections(redacted_text)

    # Stash raw Header + Contact + Links for Contact scoring. The
    # scorer runs EMAIL/PHONE/LINKEDIN regexes that don't match the
    # [EMAIL_1] / [PHONE_1] / [LINKEDIN_1] placeholders, so it needs
    # pre-redaction text to detect anything. Downstream consumers still
    # get the redacted ``sections`` dict, so PII safety is preserved.
    raw_contact = "\n".join(
        sections.get(name, "") for name in ("Header", "Contact", "Links", "Summary")
    )

    return ParsedCV(
        sections=redacted_sections,
        pii_map=pii_map,
        raw_contact=raw_contact,
    )


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("usage: python pdf_parser.py <cv.pdf|cv.docx>")
        sys.exit(1)
    result = parse_cv(sys.argv[1])
    print(json.dumps(
        {k: v[:200] + ("..." if len(v) > 200 else "") for k, v in result.items()},
        indent=2,
    ))
