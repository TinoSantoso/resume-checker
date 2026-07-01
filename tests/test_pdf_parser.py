"""Tests for app/pdf_parser.py: section segmentation + DOCX support."""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from app.pdf_parser import (
    _classify_line,
    extract_text,
    parse_cv,
    segment_sections,
)


# --------------------------------------------------------------------------- #
# Section classification (header detection)
# --------------------------------------------------------------------------- #

class TestClassifyLine:
    @pytest.mark.parametrize("line,expected", [
        ("Summary", "Summary"),
        ("Professional Summary", "Summary"),
        ("summary", "Summary"),  # case-insensitive
        ("  SUMMARY  ", "Summary"),
        ("Work Experience", "Experience"),
        ("Employment History", "Experience"),
        ("Skills", "Skills"),
        ("Technical Skills", "Skills"),
        ("Education", "Education"),
        ("Educational Background", "Education"),
        ("Contact", "Contact"),
        ("Contact Information", "Contact"),
        ("Personal Details", "Contact"),
        ("Projects", "Projects"),
        ("Certifications", "Certifications"),
        ("License & Certifications", "Certifications"),
    ])
    def test_recognized_headers(self, line, expected):
        assert _classify_line(line) == expected

    @pytest.mark.parametrize("line", [
        "Tino Santoso",
        "Senior Software Engineer with 5 years experience",
        "Led migration of monolith to microservices.",
        "Python, Go, TypeScript",
        "Bachelor of Computer Science, ITB, 2020",
    ])
    def test_unrecognized_lines_return_none(self, line):
        assert _classify_line(line) is None


# --------------------------------------------------------------------------- #
# segment_sections: pre-header goes to "Header" bucket
# --------------------------------------------------------------------------- #

class TestSegmentSections:
    def test_pre_header_content_goes_to_header_bucket(self):
        text = (
            "Tino Santoso\n"
            "tino@gmail.com | +62 812 1234 5678\n"
            "\n"
            "Summary\n"
            "Senior engineer with 8 years of experience in Python and Go.\n"
        )
        sections = segment_sections(text)
        assert "Header" in sections
        assert "tino@gmail.com" in sections["Header"]
        assert sections["Summary"].startswith("Senior engineer")

    def test_empty_input_returns_empty_dict(self):
        assert segment_sections("") == {}

    def test_only_header_no_sections(self):
        sections = segment_sections("Just a name\n+1 555 1234")
        assert "Header" in sections
        assert "Summary" not in sections

    def test_multiple_sections_ordered(self):
        text = (
            "Tino\n"
            "Summary\n"
            "A summary line here.\n"
            "Experience\n"
            "Did things.\n"
            "Skills\n"
            "Python\n"
        )
        sec = segment_sections(text)
        # All three detected, in declared order
        assert list(sec.keys())[:4] == ["Header", "Summary", "Experience", "Skills"]


# --------------------------------------------------------------------------- #
# DOCX path (added in P0.2)
# --------------------------------------------------------------------------- #

class TestDocxExtraction:
    def test_parse_cv_dispatches_to_docx(self, tmp_path):
        from docx import Document
        d = Document()
        d.add_heading("Jane Doe", 0)
        d.add_paragraph("jane@example.com | +1 555 123 4567 | linkedin.com/in/janedoe")
        d.add_heading("Professional Summary", level=1)
        d.add_paragraph(
            "Senior Data Engineer with 6+ years of experience in Python, Spark, and AWS."
        )
        d.add_heading("Skills", level=1)
        d.add_paragraph("Python, SQL, Spark, Airflow, AWS")
        path = tmp_path / "jane.docx"
        d.save(path)
        sections = parse_cv(path)
        assert "Summary" in sections
        assert "Skills" in sections
        # PII must be redacted by default (C1 backlog); use render_unmasked
        # to restore originals for the recruiter view.
        assert "jane@example.com" not in sections["Header"]
        assert "[EMAIL_" in sections["Header"]
        unmasked = sections.render_unmasked() if hasattr(sections, "render_unmasked") else sections
        assert "jane@example.com" in unmasked["Header"]
        assert "Python" in sections["Skills"]

    def test_unsupported_extension_raises_value_error(self, tmp_path):
        bad = tmp_path / "cv.txt"
        bad.write_text("not a CV")
        with pytest.raises(ValueError, match="Unsupported CV format"):
            parse_cv(bad)

    def test_pdf_still_works(self):
        # Regression: existing PDF sample must still parse.
        pdf = Path(__file__).resolve().parent.parent / "data" / "sample_strong.pdf"
        if not pdf.exists():
            pytest.skip("sample_strong.pdf not present in data/")
        sections = parse_cv(pdf)
        assert "Experience" in sections
        assert "Skills" in sections


# --------------------------------------------------------------------------- #
# DOCX end-to-end (P0.2 — task #3: ensure DOCX path is fully exercised)
# --------------------------------------------------------------------------- #

class TestDocxEndToEnd:
    """End-to-end DOCX tests against hand-built fixtures in data/.

    These are NOT auto-generated inside the test (so failures point at the
    real file, not at the test setup). Regenerate with:
        .venv/bin/python scripts/gen_docx_fixtures.py
    """

    def test_tino_shaped_docx_parses_into_known_sections(self):
        """The DOCX parser produces the major sections (Header/Contact,
        Experience, Skills, Education). The Summary section is carved
        from Header at score_cv time, not at parse_cv time — so we only
        check sections that are detected at the parse stage here."""
        docx = Path(__file__).resolve().parent.parent / "data" / "tino_shaped.docx"
        if not docx.exists():
            pytest.skip("tino_shaped.docx fixture not present — run scripts/gen_docx_fixtures.py")
        sections = parse_cv(docx)
        # Major sections must be detected at parse time
        assert "Header" in sections or "Contact" in sections
        assert "Experience" in sections
        assert "Skills" in sections
        assert "Education" in sections
        # Content must survive the docx → text round-trip
        header_blob = sections.get("Header", "") + sections.get("Contact", "")
        assert "TINO APRIKA SANTOSO" in header_blob
        assert "Laravel" in sections["Skills"]
        assert "Universitas Islam Indonesia" in sections["Education"]
        # The substantive summary paragraph is initially inside the Header
        # (the carve-out happens in score_cv, not parse_cv)
        assert "Passionate Full Stack developer" in header_blob

    def test_tino_shaped_docx_summary_carved_at_score_time(self):
        """Verify the inline summary carve-out (P1.6) works for the DOCX
        path the same way it does for PDF."""
        from app.scorer import score_cv
        docx = Path(__file__).resolve().parent.parent / "data" / "tino_shaped.docx"
        if not docx.exists():
            pytest.skip("tino_shaped.docx fixture not present")
        rep = score_cv(docx)
        # The carved Summary section should score high (50 words + role mention)
        assert rep.sections["Summary"].score >= 5.0, (
            f"Summary score {rep.sections['Summary'].score} too low — "
            "carve-out likely didn't fire for DOCX path"
        )
        # The role keyword 'developer' should be detected in the carved content
        assert any("role" in e.lower() for e in rep.sections["Summary"].evidence)

    def test_tino_shaped_docx_scores_reasonable(self):
        """DOCX version of tino_actual.pdf. Same content, different format.
        Score should be in a sensible band — we don't expect exact match
        to the PDF because text extraction differs (whitespace, formatting
        cues)."""
        from app.scorer import score_cv
        docx = Path(__file__).resolve().parent.parent / "data" / "tino_shaped.docx"
        if not docx.exists():
            pytest.skip("tino_shaped.docx fixture not present")
        rep = score_cv(docx)
        # We don't assert grade band — the heuristic is intentionally
        # looser for DOCX because text wrapping differs from PDF. We
        # DO assert the score is in a sane range and Summary carved
        # successfully.
        assert 2.0 <= rep.overall <= 7.0, f"overall {rep.overall} outside expected band"
        assert rep.sections["Summary"].score >= 5.0

    def test_strong_candidate_docx_scores_high(self):
        """A well-structured DOCX should score in the upper band."""
        from app.scorer import score_cv
        docx = Path(__file__).resolve().parent.parent / "data" / "strong_candidate.docx"
        if not docx.exists():
            pytest.skip("strong_candidate.docx fixture not present")
        rep = score_cv(docx)
        # The strong candidate has 8 years experience, 2 jobs with metrics,
        # 14 skills, 2 degrees, 2 certs. Expect C or better.
        assert rep.overall >= 6.0, f"expected strong CV to score >= 6, got {rep.overall}"
        # Summary should be high (8 years + role + appropriate length)
        assert rep.sections["Summary"].score >= 8.0, (
            f"Summary should score high for a strong CV, got {rep.sections['Summary'].score}"
        )
        # Experience should have bullets with metrics
        assert rep.sections["Experience"].score >= 6.0, (
            f"Experience should score well, got {rep.sections['Experience'].score}"
        )

    def test_docx_with_table_renders_cells(self):
        """Tables in a DOCX should be rendered row-by-row with cells
        joined by ' | ' (matches the test_pdf_parser.py table test)."""
        from app.pdf_parser import _extract_docx_text
        from docx import Document
        d = Document()
        d.add_paragraph("Header row")
        table = d.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Skill"
        table.rows[0].cells[1].text = "Level"
        table.rows[1].cells[0].text = "Python"
        table.rows[1].cells[1].text = "Advanced"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            d.save(tf.name)
            text = _extract_docx_text(tf.name)
        # Header row + table row, joined with " | "
        assert "Skill | Level" in text
        assert "Python | Advanced" in text

    def test_docx_list_bullet_style_gets_bullet_prefix(self):
        """Paragraphs with List Bullet style must be prefixed with '• ' so
        the experience heuristic recognizes them as bullets. Without this
        fix, real Word docs (which use List Bullet by default) would be
        mis-classified as paragraph format and penalized."""
        from app.pdf_parser import _extract_docx_text
        from docx import Document
        d = Document()
        d.add_paragraph("Senior Engineer  BigTech Co  Jan 2021 - Present")
        d.add_paragraph(
            "Architected pipeline processing 50K events/second.",
            style="List Bullet",
        )
        d.add_paragraph(
            "Mentored 4 engineers; 2 promoted to Senior.",
            style="List Bullet",
        )
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            d.save(tf.name)
            text = _extract_docx_text(tf.name)
        # Each bullet paragraph must start with the bullet marker
        assert "• Architected pipeline" in text
        assert "• Mentored 4 engineers" in text
        # The role-header line (not bullet-styled) must NOT have a bullet
        # marker prepended
        for line in text.splitlines():
            if "BigTech Co" in line:
                assert not line.startswith("• "), (
                    f"role header should not be bulleted: {line!r}"
                )

    def test_docx_list_number_style_passes_through(self):
        """Numbered list items should not get a • prefix (the experience
        heuristic's \\d+[.)] pattern catches numbered items)."""
        from app.pdf_parser import _extract_docx_text
        from docx import Document
        d = Document()
        d.add_paragraph("First step", style="List Number")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            d.save(tf.name)
            text = _extract_docx_text(tf.name)
        assert "First step" in text
        assert "• First step" not in text

    def test_docx_bullets_produce_high_experience_score(self):
        """End-to-end: a real Word-style DOCX with List Bullet paragraphs
        should NOT be flagged as paragraph format and should score well
        on Experience."""
        from app.scorer import score_cv
        from docx import Document
        d = Document()
        d.add_heading("Jane Doe", 0)
        d.add_paragraph("jane@example.com | +1 555 010 2024")
        d.add_heading("Experience", level=1)
        d.add_paragraph("Senior Engineer  BigTech  2020 - Present")
        for bullet in [
            "Shipped pipeline processing 50K events/second with 99.99% uptime.",
            "Led migration to Kubernetes reducing cost by 30%.",
            "Mentored 4 engineers; 2 promoted to Senior within a year.",
            "Built payment integration handling $120M/year revenue.",
        ]:
            d.add_paragraph(bullet, style="List Bullet")
        d.add_heading("Skills", level=1)
        d.add_paragraph("Python, Go, Kubernetes, AWS, PostgreSQL")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            d.save(tf.name)
            rep = score_cv(tf.name)
        # The paragraph-format penalty should NOT have fired
        issues_text = " ".join(rep.sections["Experience"].issues)
        assert "paragraph format" not in issues_text.lower(), (
            f"List Bullet paragraphs should not be flagged as paragraph format: {issues_text}"
        )
        assert rep.sections["Experience"].score >= 7.0, (
            f"Expected high Experience score with bulleted metrics, got {rep.sections['Experience'].score}"
        )


# --------------------------------------------------------------------------- #
# PII redaction integration (C1 backlog)
# --------------------------------------------------------------------------- #

class TestPIIRedactionIntegration:
    """End-to-end: parse_cv() must redact PII before downstream use.

    Backward compat: callers that iterated the old dict return value
    still work because ParsedCV.__iter__ yields (name, text) tuples.
    """

    def _write_docx(self, tmp_path, body_lines):
        from docx import Document
        docx_path = tmp_path / "cv.docx"
        doc = Document()
        for line in body_lines:
            doc.add_paragraph(line)
        doc.save(str(docx_path))
        return docx_path

    def test_parse_cv_returns_parsedcv_with_pii_redacted(self, tmp_path):
        docx_path = self._write_docx(tmp_path, [
            "Contact",
            "Email: tino.santoso@gmail.com",
            "Phone: 0812-3456-7890",
            "LinkedIn: linkedin.com/in/tino",
            "Summary",
            "Senior engineer with 10 years experience.",
        ])
        result = parse_cv(docx_path)
        # Backward-compat: still iterable as (name, text) pairs
        assert hasattr(result, "items")
        assert "Contact" in result
        # PII must be redacted
        contact_text = result["Contact"]
        assert "tino.santoso@gmail.com" not in contact_text
        assert "0812-3456-7890" not in contact_text
        assert "linkedin.com/in/tino" not in contact_text
        # Placeholders present
        assert "[EMAIL_1]" in contact_text
        assert "[PHONE_1]" in contact_text
        assert "[LINKEDIN_1]" in contact_text
        # pii_map round-trip restores originals
        assert result.pii_map["[EMAIL_1]"] == "tino.santoso@gmail.com"

    def test_render_unmasked_restores_original_pii(self, tmp_path):
        docx_path = self._write_docx(tmp_path, [
            "Contact",
            "Email: a@b.com",
            "Summary",
            "Solid background.",
        ])
        result = parse_cv(docx_path)
        unmasked = result.render_unmasked()
        assert "a@b.com" in unmasked["Contact"]

    def test_redact_pii_false_keeps_raw_text(self, tmp_path):
        docx_path = self._write_docx(tmp_path, [
            "Contact",
            "Email: a@b.com",
            "Summary",
            "Hi.",
        ])
        result = parse_cv(docx_path, redact_pii=False)
        assert "a@b.com" in result["Contact"]
        assert result.pii_map == {}

    def test_pii_in_summary_section_also_redacted(self, tmp_path):
        docx_path = self._write_docx(tmp_path, [
            "Contact",
            "Some text",
            "Summary",
            "Reach me at tino@gmail.com or 0812-3456-7890 for opportunities.",
        ])
        result = parse_cv(docx_path)
        summary = result["Summary"]
        assert "tino@gmail.com" not in summary
        assert "0812-3456-7890" not in summary
        assert "[EMAIL_" in summary
        assert "[PHONE_" in summary

    def test_dict_style_access_still_works(self, tmp_path):
        """Existing code that does parse_cv(path)['Skills'] must not break."""
        docx_path = self._write_docx(tmp_path, [
            "Skills",
            "Python, AWS, Docker",
        ])
        result = parse_cv(docx_path)
        assert "Python" in result["Skills"]


# --------------------------------------------------------------------------- #
# Regression tests: segment_sections wrapped-paragraph bug (audit fix-now)
# --------------------------------------------------------------------------- #

class TestSegmentSectionsWrappedParagraph:
    """Audit (docs/summary_carveout_audit.md): on some CVs the Summary
    paragraph is physically located in the PDF reading order AFTER another
    section header (typically Education), so a line-by-line segmenter
    buckets it into the wrong section. The 'Summary' header at the top
    of the doc then only captures the trailing wrapped line.

    These tests pin down the desired behavior: a Summary header should
    gather ALL Summary-like prose, even if it appears later in the doc.
    """

    def test_segment_sections_preserves_multiline_summary_paragraph(self):
        """Real bleed pattern: Summary header appears at top, but the
        Summary paragraph ALSO appears again later in the doc (after
        Education). All prose lines of the Summary paragraph must land
        in the Summary bucket — none in Education."""
        text = (
            "Tino Santoso\n"
            "tino@gmail.com | +62 812 1234 5678\n"
            "\n"
            "Summary\n"
            "products in the last 2 years.\n"
            "Experience\n"
            "Senior PM at Acme  Jan 2020 - Present\n"
            "Did things.\n"
            "Skills\n"
            "Python, SQL\n"
            "Education\n"
            "B.Sc. CS — UI, 2018\n"
            "Achieved 18% revenue growth via data-driven pricing experiments.\n"
            "Launched 3 A/B tests across the funnel that lifted conversion by 12%.\n"
            "Led cross-functional team of 5 engineers and 2 designers to ship in 6 weeks.\n"
        )
        sections = segment_sections(text)
        assert "Summary" in sections
        # The full Summary paragraph must be in Summary even though
        # Education header appears BEFORE those prose lines in reading order.
        assert "revenue growth" in sections["Summary"]
        assert "A/B tests" in sections["Summary"]
        assert "ship in 6 weeks" in sections["Summary"]
        # The Summary paragraph must NOT have been stranded in Education.
        assert "revenue growth" not in sections.get("Education", "")

    def test_segment_sections_no_bleed_when_summary_first_section(self):
        """Summary header first; the real Summary paragraph appears later
        in the doc (after Education). Verify Summary bucket recovers the
        full paragraph and Education bucket is NOT polluted."""
        text = (
            "Summary\n"
            "products in the last 2 years.\n"
            "Experience\n"
            "Senior Frontend Engineer\n"
            "• Built the seller onboarding flow; conversion 32%\n"
            "Skills\n"
            "React\n"
            "Education\n"
            "B.Sc. CS — UI, 2018\n"
            "rina.anggraini@gmail.com\n"
            "Senior Frontend Engineer with 6 years of experience building customer-facing\n"
            "web apps. Specialised in React, TypeScript, and design systems. Led the design\n"
            "system rebuild at Bukalapak; shipped 4 net-new\n"
            "React, TypeScript, Next.js, GraphQL, Storybook\n"
        )
        sections = segment_sections(text)
        assert "Summary" in sections
        summary_body = sections["Summary"]
        # The full Summary paragraph must be in Summary, not in Education.
        assert "6 years of experience" in summary_body
        assert "design systems" in summary_body
        # Joined across line wraps; collapse newlines for assertion.
        joined = " ".join(summary_body.split())
        assert "design system rebuild" in joined
        # Education must NOT contain Summary prose.
        assert "6 years of experience" not in sections.get("Education", "")

    def test_segment_sections_v_pm_metrics_driven_regression(self):
        """Regression against the real bleeding PDF. After the fix, the
        Summary section must contain the full paragraph (≥200 chars, ≥3
        lines) and a substantive keyword, not just the trailing wrapped
        line '(+18%) and retention (+12 NPS).' which the current code
        catches alone."""
        pdf = (
            Path(__file__).resolve().parent.parent
            / "data" / "validation" / "v_pm_metrics_driven.pdf"
        )
        if not pdf.exists():
            pytest.skip("v_pm_metrics_driven.pdf not present in data/validation/")
        raw = extract_text(pdf)
        sections = segment_sections(raw)
        assert "Summary" in sections, "Summary section must be detected"
        summary = sections["Summary"]
        # Real Summary paragraph is ~600 chars; pre-fix this is ~30 chars.
        assert len(summary) >= 200, (
            f"Summary too short ({len(summary)} chars) — wrapped-paragraph "
            f"bug regressed. Got: {summary!r}"
        )
        assert len([ln for ln in summary.splitlines() if ln.strip()]) >= 3, (
            f"Summary must have >= 3 lines, got {summary!r}"
        )
        # The real Summary mentions growth / pricing / revenue.
        lowered = summary.lower()
        assert any(
            kw in lowered
            for kw in ("revenue", "growth", "pricing", "launched", "ship")
        ), f"Summary missing expected keyword. Got: {summary!r}"
