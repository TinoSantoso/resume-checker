"""Shared fixtures: build small fake CV text snippets and DOCX files."""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

# Skip the whole module if python-docx is not installed.
docx = pytest.importorskip("docx")

from docx import Document  # type: ignore  # noqa: E402


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """Factory fixture: build a DOCX from a list of (kind, text) lines.

    kind is one of:
      - "h"   -> heading 1 (treated as section header by the parser)
      - "p"   -> normal paragraph
      - "t"   -> table cell content (auto-spread into one row)
    """
    def _make(rows: list[tuple[str, str]], name: str = "cv.docx") -> Path:
        d = Document()
        for kind, text in rows:
            if kind == "h":
                d.add_heading(text, level=1)
            elif kind == "p":
                d.add_paragraph(text)
            elif kind == "t":
                # 3-col table, 1 row
                tbl = d.add_table(rows=1, cols=3)
                cells = text.split("|")
                for i, c in enumerate(cells[:3]):
                    tbl.rows[0].cells[i].text = c.strip()
            else:
                raise ValueError(f"unknown row kind: {kind!r}")
        out = tmp_path / name
        d.save(out)
        return out
    return _make
